# ATENÇÃO: Este arquivo foi adaptado para portfólio. Todas as variáveis de ambiente e dados sensíveis devem ser fictícios.
# Não inclua credenciais reais, senhas, nomes de empresas ou pessoas reais neste código.
import requests
import pandas as pd
from docx import Document
from bs4 import BeautifulSoup
import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import psycopg2
import os
import base64
import chardet
from dotenv import load_dotenv
import tempfile

load_dotenv()

# Constantes
API_KEY = os.getenv("API_KEY")
URL = os.getenv("URL")
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
HASH = os.getenv("COD_HASH")

#TRANSCRICAO_PATH = os.path.join(CURRENT_DIR, 'transcricao.txt')

def connect_db():
    try:
        # Conectar ao banco de dados
        conn = psycopg2.connect(
            dbname = os.getenv("DB_NAME"),
            user = os.getenv("USER"),
            password = os.getenv("PASSWORD"),
            host = os.getenv("HOST"),
            port = os.getenv("PORT") 
        )
        print("Conexão ao banco de dados estabelecida com sucesso.")
        return conn
    except psycopg2.Error as e:
        print(f"Erro ao conectar ao banco de dados: {e}")
        return None

def gerar_html_ata(transcricao_path):
    try:
        print("Tentando conectar ao banco de dados...")  # Mensagem de depuração
        engine = connect_db()
        if engine is None:
            print("Erro: Conexão ao banco de dados falhou.")
            return None

        print("Conexão ao banco de dados estabelecida com sucesso.")  # Mensagem de depuração

        print("Executando consulta SQL...")  # Mensagem de depuração
        query = "SELECT topico, contexto FROM tbl_topicos;"

        try:
            df = pd.read_sql(query, engine)
            print("Consulta SQL executada com sucesso.")  # Mensagem de depuração
        except Exception as e:
            print(f"Erro ao executar consulta SQL: {e}")
            return None

        # Exibir as primeiras linhas do DataFrame
        print("Exibindo as primeiras linhas do DataFrame:")
        print(df.head())

        def detectar_codificacao(arquivo):
            with open(arquivo, 'rb') as f:
                resultado = chardet.detect(f.read())
            return resultado['encoding']

        print("Tentando ler o arquivo de transcrição...")  # Mensagem de depuração
        with open(transcricao_path, 'r', encoding=detectar_codificacao(transcricao_path)) as file:
            transcricao = file.read()
        print("Arquivo de transcrição lido com sucesso.")  # Mensagem de depuração
    
    except Exception as e:
        raise RuntimeError(f"Erro ao processar o arquivo: {e}")

        

    except Exception as e:
        print(f"Erro ao ler tabela: {e}")
        return None


    html_resultado = ""

    df.rename(columns={
        'Tópico ': 'topico',
        'Contexto': 'contexto'
    }, inplace=True)

    for _, row in df.iterrows():
        topico = row['topico']
        contexto = row['contexto']
        
        resposta = chamar_api(transcricao, contexto)
        resposta_limpa = limpar_texto(resposta)
        
        html_resultado += f"<h2>{""}</h2>\n"
        html_resultado += resposta_limpa
        html_resultado += "<hr>\n"

    return html_resultado

def chamar_api(str_reuniao, str_topico):
    data = {
        "inputs": {
            "str_reuniao": str_reuniao,
            "str_topico": str_topico,
        }
    }
    
    try:
        response = requests.post(f"{URL}/api/templates/{HASH}/execute", json=data, headers={"X-Api-Key": API_KEY})
        response.raise_for_status()
        return response.text
    except requests.exceptions.RequestException as e:
        return f"Erro: {e}"

def limpar_texto(texto):
    return re.sub(r"```html|```|`", "", texto)


def tratar_html_no_word(html_content):
    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.find_all(['h1', 'h2', 'p', 'table']):
        if element.name == 'h1':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(178, 34, 34)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif element.name == 'h2':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(14)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif element.name == 'p':
            paragraph = doc.add_paragraph(element.get_text())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(12)
        
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                cols = len(rows[0].find_all(['th', 'td']))
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                
                for row_idx, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for col_idx, cell in enumerate(cells):
                        cell_text = cell.get_text()
                        table.cell(row_idx, col_idx).text = cell_text

                        if row_idx == 0:
                            table.cell(row_idx, col_idx).paragraphs[0].runs[0].bold = True
                            table.cell(row_idx, col_idx).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

def substituir_no_word(modelo_docx_binario, html_content):
    doc = Document(BytesIO(modelo_docx_binario))
    
    for paragraph in doc.paragraphs:
        if '@ata' in paragraph.text:
            paragraph.text = paragraph.text.replace('@ata', html_content)
    
    return doc

def converter_arquivo_para_base64(caminho_arquivo):
    with open(caminho_arquivo, "rb") as arquivo:
        arquivo_binario = arquivo.read()
        return base64.b64encode(arquivo_binario).decode('utf-8')

def main(transcricao_path):
    print("Gerando HTML da ata...")
    html_content = gerar_html_ata(transcricao_path)
    if not html_content:
        print("Erro ao gerar HTML da ata.")
        return
    
#    modelo_docx_binario = get_file_from_db(2)  # ID 2 para o arquivo Word
    
    print("Tratando HTML e gerando documento Word...")
    doc = tratar_html_no_word(html_content)
    
    
    # Usar um diretório temporário para o arquivo gerado
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_file_name = temp_file.name  # Nome do arquivo temporário
        doc.save(temp_file_name)  # Salva o documento no arquivo temporário
        print(f"Documento Word gerado temporariamente em: {temp_file_name}")
    
    # Converter o arquivo salvo em base64
    print("Convertendo documento para base64...")
    doc_base64 = converter_arquivo_para_base64(temp_file_name)
    print("Documento convertido para base64 com sucesso!")
    print(doc_base64)
    return doc_base64


if __name__ == "__main__":
    main()
