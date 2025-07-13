import io
import pandas as pd
import requests
import json
import os
from docx import Document
import numpy as np
import re
from bs4 import BeautifulSoup
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tqdm import tqdm
import psycopg2
import zipfile
from xml.etree.ElementTree import ElementTree, fromstring
from datetime import datetime
import PyPDF2
import unidecode

# ====== INÍCIO DA ANONIMIZAÇÃO DE INFORMAÇÕES SENSÍVEIS ======
# As informações abaixo foram alteradas para fins fictícios e de portfólio.
# Nenhuma credencial, senha, chave de API ou caminho real está presente neste código.
# Caso utilize este código, substitua pelos seus próprios dados.

# Exemplo de chave de API fictícia
apiKey = "API_KEY_FICTICIA_PARA_PORTFOLIO"

def scope_item_resumo(apiKey, str_scope_item, pergunta, contexto):
    str_idioma = "Portugues/BR"
    str_topico = f'{pergunta} {contexto}'

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_scope_item": str_scope_item,
            "str_topico": str_topico,
        }
    }

    response = requests.post(f"{url}/api/templates/66578b1b5d54205ab618cd85/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def scope_item_descricao(apiKey, str_scope_item):
    str_idioma = "Portugues/BR"

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_scope_item": str_scope_item,
        }
    }

    response = requests.post(f"{url}/api/templates/6661ca2d4c565b1488fb4ccb/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def read_docx_text_xml(file_path):
    
    # file_path='/Users/rbarrionuevo/Downloads/BPD_Potencial/planilha_BPD001 - EO - Estrutura Organizacional Financas/Scope_Item/1QM01_S4HANA2023_BPD_EN_BR pt-BR - Gestão Avançada de Crédito.docx'
    
    try:
        # Abre o arquivo .docx como um arquivo ZIP
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            # Abre o XML que contém o texto principal do documento
            with zip_ref.open('word/document.xml') as xml_file:
                # Lê o conteúdo do XML
                xml_content = xml_file.read()

                # Parseia o XML
                xml_tree = fromstring(xml_content)

                # Define o namespace usado no XML do Word
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                # Extrai o texto do XML
                text_builder = []
                for t in xml_tree.iterfind('.//w:t', namespaces=ns):
                    if t.text is not None:
                        text_builder.append(t.text)

                # Retorna o texto concatenado
                return ''.join(text_builder)

    except Exception as e:
        print(f"Erro ao ler o arquivo .docx: {e}")

def save_text_to_file(text, output_file):
    try:
        # Abre o arquivo de saída para escrita
        with open(output_file, 'w', encoding='utf-8') as f:
            # Escreve o texto no arquivo
            f.write(text)
            print(f"Texto salvo com sucesso em {output_file}")
    except Exception as e:
        print(f"Erro ao salvar o arquivo: {e}")
        
def replace_especial(texto):
    # Definindo os caracteres especiais a serem substituídos por _
    padrao = r'[\/,.\|* ?~`!@#$%^&*()_+=\-\[\]{}<>:;\'\"\\]'
    
    # Substituindo os caracteres especiais por _
    texto_substituido = re.sub(padrao, '_', texto)
    
    # Substituindo duplicidade de underscores por um único underscore
    texto_substituido = re.sub(r'_+', '_', texto_substituido)
    
    # Removendo underscore no início ou final da string, se houver
    texto_substituido = texto_substituido.strip('_')
    
    return texto_substituido
        

def create_connection():
    try:
        # Defina as informações de conexão
        connection = psycopg2.connect(
            dbname="db_ScopeItems",
            user="postgres",
            password="Stefanini@2024",
            host="ec2-44-203-195-255.compute-1.amazonaws.com",
            port="2203"
        )

        # Crie um cursor
        cursor = connection.cursor()        
        return connection, cursor
    except (Exception, psycopg2.DatabaseError) as error:
        print(f"Erro ao conectar ao banco de dados: {error}")
        return None, None

def close_connection(connection, cursor):
    if cursor:
        cursor.close()
    if connection:
        connection.close()
    print("Conexão com o banco de dados fechada.")

def padronizar_formato_data(data):
    try:
        # Tentar analisar a data nos formatos possíveis
        data_formatada = datetime.strptime(data, "%d-%m-%Y").strftime("%Y-%m-%d")
    except ValueError:
        try:
            data_formatada = datetime.strptime(data, "%Y-%m-%d").strftime("%Y-%m-%d")
        except ValueError:
            try:
                data_formatada = datetime.strptime(data, "%d-%m-%y").strftime("%Y-%m-%d")
            except ValueError:
                print(f"Formato de data não reconhecido: {data}")
                return None

    return data_formatada

def Insert_Scope_Item(diretorio):
    
    #diretorio = caminho_scope_itens
    
    try:
        # Percorrer os arquivos no diretório especificado
        for arquivo in os.listdir(diretorio):
            print(arquivo)
            caminho_completo = os.path.join(diretorio, arquivo)
            if arquivo.endswith(".docx"): 
                try:
                    # caminho_completo = '/Users/rbarrionuevo/Downloads/BPD_Potencial/planilha_BPD001 - EO - Estrutura Organizacional Financas/Scope_Item/1QM01_S4HANA2023_BPD_EN_BR pt-BR - Gestão Avançada de Crédito.docx'
                    
                    conteudo_documento = read_docx_text_xml(caminho_completo)
                    
                    # Salva o texto extraído em um arquivo .txt no mesmo diretório e com nome baseado no arquivo original
                    nome_arquivo_txt = os.path.splitext(arquivo)[0] + ".txt"
                    output_file = os.path.join(diretorio, nome_arquivo_txt)
                    save_text_to_file(conteudo_documento, output_file)
                    
                except Exception as e:
                    print(f"Erro ao processar o arquivo {caminho_completo}: {e}")
                    continue

                descricao = scope_item_descricao(apiKey, conteudo_documento).replace("```", "").replace("json", "")
                descricao_itens = json.loads(descricao)
                
                codigo_arquivo = arquivo.split('_')
                
                # Busca cabeçalho do documento
                codigo = codigo_arquivo[0]
                nome = descricao_itens.get("Nome")
                versao = descricao_itens.get("Versão")
                idioma = descricao_itens.get("Idioma")
                componente_sap = descricao_itens.get("Componente SAP")
                data_documento = descricao_itens.get("Data do Documento")
                data_documento = padronizar_formato_data(data_documento)
                
                conn, cur = create_connection()
                
                if conn is not None and cur is not None:
                    try:
                        # Verificar se o registro já existe
                        cur.execute("SELECT COUNT(*) FROM tbl_scopeitens WHERE codigo = %s", (codigo,))
                        exists = cur.fetchone()[0]
                        if exists:
                            print(f"O registro com código {codigo} já existe. Pulando para o próximo arquivo.")
                            continue
                        
                        # Execute a query para selecionar todos os registros da tabela
                        cur.execute("SELECT nIdPK, topico, contexto FROM tbl_perguntasScope;")
                        
                        # Obtenha todos os registros
                        records = cur.fetchall()
                        
                        # Itere sobre os registros e faça algo com eles
                        for record in records:
                            nIdPK, topico, contexto = record
                            
                            # Verificar se o registro específico já existe
                            cur.execute("SELECT COUNT(*) FROM tbl_scopeitens WHERE nomedocumento = %s AND topico = %s", (arquivo, topico))
                            specific_exists = cur.fetchone()[0]
                            if specific_exists:
                                print(f"O registro com nomedocumento {arquivo} e topico {topico} já existe. Pulando para o próximo registro.")
                                continue
                            
                            scope_item_resumo_result = scope_item_resumo(apiKey, conteudo_documento, topico, contexto)
                            scope_item_resumo_result = scope_item_resumo_result.replace("```", "").replace("json", "").rstrip()
                            
                            insert_query = '''
                            INSERT INTO tbl_scopeitens (codigo, nomedocumento, topico, contexto, resposta, idioma, componente_sap, data_documento)
                            VALUES (%(codigo)s, %(arquivo)s, %(topico)s, %(contexto)s, %(scope_item_resumo_result)s, %(idioma)s, %(componente_sap)s, %(data_documento)s)
                            '''
                            
                            params = {
                                'codigo': codigo,
                                'arquivo': arquivo,
                                'topico': topico,
                                'contexto': contexto,
                                'scope_item_resumo_result': scope_item_resumo_result,
                                'idioma': idioma,
                                'componente_sap': componente_sap,
                                'data_documento': data_documento
                            }
                            
                            cur.execute(insert_query, params)
                            conn.commit()

                    except (Exception, psycopg2.DatabaseError) as error:
                        print(f"Erro ao buscar registros: {error}")
                    
                    # Feche a conexão após o uso
                    close_connection(conn, cur)
                
    except Exception as e:
        print(f'Erro ao processar o arquivo {caminho_completo}: {e}')


def busca_scopeItem(codigo_scopeitem):
    
    def create_connection():
        try:
            # Defina as informações de conexão
            connection = psycopg2.connect(
                dbname="db_ScopeItems",
                user="postgres",
                password="Stefanini@2024",
                host="ec2-44-203-195-255.compute-1.amazonaws.com",
                port="2203"
            )

            # Crie um cursor
            cursor = connection.cursor()        
            return connection, cursor
        except (Exception, psycopg2.DatabaseError) as error:
            print(f"Erro ao conectar ao banco de dados: {error}")
            return None, None

    def close_connection(connection, cursor):
        if cursor:
            cursor.close()
        if connection:
            connection.close()
            print("Conexão com o banco de dados fechada.")

    def busca_resumo_scopeItem(codigo):
               
        print(codigo)
        
        conn, cur = create_connection()
        if conn is None or cur is None:
            return None
        
        try:
            cur.execute("SELECT resposta FROM vw_Busca_Scopeitem2 WHERE codigo = %s", (codigo,))
            row = cur.fetchone()  # Usar fetchone() para obter um único resultado
            if row:
                # Aqui row[0] é o valor retornado do banco de dados
                return json.dumps(row[0])  # Retorna a lista de dicionários convertida em JSON
            else:
                return None
        except Exception as e:
            print(f"Erro ao buscar dados: {e}")
            return None
        finally:
            close_connection(conn, cur)
        
    return busca_resumo_scopeItem(codigo_scopeitem)

def cria_df_BPH(caminho_base_bpd, nome_arquivo_xlsx, caminho_scope_itens):
    
    Insert_Scope_Item(caminho_scope_itens)
    
    caminho_arquivo = f'{caminho_base_bpd}{nome_arquivo_xlsx}'

    # Ler o arquivo Excel
    df_bph = pd.read_excel(caminho_arquivo)

    # Preencher valores ausentes na coluna 'nome_bpd'
    df_bph['nome_bpd'] = df_bph['nome_bpd'].fillna('')

    # Inicializar o DataFrame para armazenar o conteúdo
    resultado_df = pd.DataFrame(columns=['nome_bpd', 'scope_item', 'cod_gap', 'scope_items_path', 'scope_items_docx'])

    # Agrupar 'scope_item' e 'cod_gap' por 'processo' e remover duplicados
    grupo_processo_scope_item_gap = df_bph.groupby('nome_bpd').agg({'scope_item': lambda x: list(set(x)), 'cod_gap': lambda x: list(set(x))})

    # Exibir os dados agrupados e buscar os resumos dos Scope Itens
    for nome_bpd, (scope_itens, cod_gaps) in grupo_processo_scope_item_gap.iterrows():
        for scope_item in scope_itens:
            if pd.notnull(scope_item) and isinstance(scope_item, str):
                # Buscar o resumo do scope_item no banco de dados
                print(scope_item)
                resumo_scopeItem = busca_scopeItem(scope_item)
                
                if resumo_scopeItem:
                    # Atualizar df_bph com os dados encontrados
                    df_bph.loc[(df_bph['scope_item'] == scope_item) & (df_bph['nome_bpd'] == nome_bpd), 'scope_items_path'] = ''  # Não usado mais
                    df_bph.loc[(df_bph['scope_item'] == scope_item) & (df_bph['nome_bpd'] == nome_bpd), 'scope_items_docx'] = resumo_scopeItem
                    
                    # Atualizar o resultado_df (opcional, se necessário)
                    resultado_df = pd.concat([resultado_df, pd.DataFrame({
                        'nome_bpd': [nome_bpd],
                        'scope_item': [scope_item],
                        'cod_gap': [cod_gaps],
                        'scope_items_path': [''],  # não usa mais
                        'scope_items_docx': [resumo_scopeItem]
                    })], ignore_index=True)
                else:
                    print(f"Nenhum resumo encontrado para o scope item: {scope_item} do processo: {nome_bpd}")
            else:
                print(f"Scope item inválido encontrado: {scope_item} para o processo: {nome_bpd}")

    return df_bph
                        
#%% busca_variaveis
def busca_variaveis(df_bph, nome_bpd, apiKey, caminho_base_bpd, modulo):
 
    
    def gerar_h1(texto, fonte, tamanho, cor):
        """
        Gera uma tag <h1> com estilo personalizado.
        
        Args:
        texto (str): O texto a ser exibido dentro da tag <h1>.
        fonte (str): A fonte a ser usada no texto.
        tamanho (str): O tamanho da fonte.
        cor (str): A cor do texto.
        
        Returns:
        str: Uma string contendo o HTML formatado para o <h1>.
        """
        html = f"""<!DOCTYPE html>
    <html lang="pt-BR">
    <head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{texto}</title>
    </head>
    <body>
    <h1><span style='color: {cor}; font-family:"{fonte}"; font-size: {tamanho}px'><b>{texto}</b></span></h1>
    </body>
    </html>"""
        return html
    
      
    
    def docx_to_txt_caminho(wCaminho):
        conteudos = {}  # Dicionário para armazenar o conteúdo de cada arquivo como texto
        try: 
            # Percorrer os arquivos no diretório especificado
            for arquivo in os.listdir(wCaminho):
                caminho_completo = os.path.join(wCaminho, arquivo)
                documento = Document(caminho_completo)
                texto = [paragrafo.text for paragrafo in documento.paragraphs]  # Extrai o texto de cada parágrafo
                conteudo_texto = '\n'.join(texto)  # Junta todos os parágrafos em uma única string
                conteudos[arquivo] = conteudo_texto  # Armazena o texto no dicionário com o nome do arquivo como chave
        except Exception as e:
            print(f"Erro ao processar arquivo: {e}")
            
        return conteudos 
    
    
    def pdf_to_txt_diretorio(caminho_pdf):
           # Abrindo o arquivo PDF em modo leitura binária
           
        texto_extraido = ''
        
        try:
            # Abrindo o arquivo PDF em modo leitura binária
            with open(caminho_pdf, 'rb') as arquivo_pdf:
                # Criando um objeto PDF reader
                leitor_pdf = PyPDF2.PdfReader(arquivo_pdf)
                # Iterando sobre cada página do PDF
                for pagina in range(len(leitor_pdf.pages)):
                    # Extraindo o texto da página atual
                    texto_extraido += leitor_pdf.pages[pagina].extract_text()
                    
        except Exception as e:
            print(f"Erro ao processar o arquivo PDF: {e}")
        
        return texto_extraido
    
    def docx_to_txt_diretorio(diretorio):
        conteudos = {}  # Dicionário para armazenar o conteúdo de cada arquivo como texto
        try: 
            # Percorrer os arquivos no diretório especificado
            for arquivo in os.listdir(diretorio):
                if arquivo.endswith(".docx"):  # Verifica se o arquivo é um docx
                    caminho_completo = os.path.join(diretorio, arquivo)
                    documento = Document(caminho_completo)
                    texto = [paragrafo.text for paragrafo in documento.paragraphs]  # Extrai o texto de cada parágrafo
                    conteudo_texto = '\n'.join(texto)  # Junta todos os parágrafos em uma única string
                    conteudos[arquivo] = conteudo_texto  # Armazena o texto no dicionário com o nome do arquivo como chave
                    
                elif arquivo.endswith(".pdf"):
                    caminho_completo = os.path.join(diretorio, arquivo)
                    conteudo_texto = pdf_to_txt_diretorio(caminho_completo)
                    conteudos[arquivo] = conteudo_texto
                    
                    
        except Exception as e:
            
            print(f"Erro ao processar diretório: {e}")
            
        return conteudos    

    
    def transicao_reuniao(diretorio, apiKey, nome_bpd, str_processo):
        resultados = {}  # Dicionário para armazenar os resultados das transcrições por tópico
        
        try:
            # Primeiro conjunto de dados (df1)
            data1 = {
                "topicos": [
                    "Best Practice",
                    "Mudanças na Best Practice",
                    "Descrição do Processo de Negócio",
                    "Localização onde este Processo de Negócio é Executado",
                    "Decisões Operacionais ou Lógica Envolvida no Processo",
                    "Considerações Legais e Políticas Específicas do Cliente",
                    "Ponto de Integração com os demais módulos",
                    "Estrutura Organizacional Envolvida",
                    "Dados Mestres Relacionados",
                    "Configurações de Sistemas Envolvidas (Deltas)",
                    "Transações e Aplicativos Envolvidos",
                    "DESENVOLVIMENTOS ENVOLVIDOS",
                    "AUTORIZAÇÕES (ROLES)",
                    "IMPACTOS ORGANIZACIONAIS"
                ],
                "contexto": [
                    "(Traga uma lista simples com o código do scope item citado e uma descrição breve sobre o documento. Se atenha somente ao código de Scope Item, e nunca códigos não relacionados a este tópico.)",
                    "(Traga o que foi dito sobre mudanças de Best Practices, quando citadas)",
                    "(Traga o que foi dito sobre a descrição dos processos de negócio, principalmente informações relacionadas a: Dados de Ativação; Entradas; Saídas; Volumetria; Frequência; Proprietários e demais complementos relevantes sobre a descrição do processo de negócio.)",
                    "(Traga o que foi dito sobre as localizações geográficas, escritórios, setores, fábricas, armazéns, etc., onde o processo é executado.)",
                    "(Traga informações sobre decisões operacionais e lógicas relacionadas a esse processo.)",
                    "(Traga informações sobre as políticas da empresa e a considerações de compliance, como SOX, GDPR, Segregation of Duty, exigências de relatórios, dentre outras políticas cruciais para o projeto e processo de negócio.)",
                    "(Traga informações sobre pontos de integração/problemas com outros módulos SAP ou componentes, etc.)",
                    "(Traga tudo o que foi dito sobre Estrutura Organizacional. Por exemplo, itens da estrutura da organização SAP necessários para dar suporte a esse processo.)",
                    "(Traga tudo o que foi dito especificamente sobre Dados Mestre)",
                    "(Traga o que foi dito sobre as configurações necessárias para os sistemas envolvidos no processo e demais informações relevantes para este tópico.)",
                    "(Traga o que foi dito sobre as transações FIORI que serão utilizadas, trazendo sempre seus nomes e, quando citado, seus códigos ou IDs e demais informações relevantes.)",
                    "(Traga o que foi dito sobre os desenvolvimentos (gaps) identificados para este processo, incluindo informações relevantes sobre estes.)",
                    "(Traga o que foi dito sobre quais serão as práticas de autorizações, as roles, objeto de dados, motivos, dentre outras informações relevantes para este tópico.)",
                    "(Traga o que foi dito sobre quais são os impactos organizacionais percebidos para este processo, incluindo o nome dos setores impactados, descrições e demais informações relevantes.)"
                ]
            }
    
            # Segundo conjunto de dados (df2)
            data2 = {
                "topicos": [
                    "Requerimento do Processo de Negócio",
                    "Detalhamento dos Passos do Processo"
                ],
                "contexto": [
                    "(Traga o que foi dito sobre o que é o processo de [str_processo] e, quando houver, os requerimentos do processo de negócio e seus devidos subprocessos.)",
                    "(Traga o que foi dito sobre todas as etapas gerais do processo. Busque informações como: Etapa do Processo; Ref. para Requisitos Aplicáveis; Considerações para Entrega; Descrição de GAPs; Descrição de Solução.)"
                ]
            }
    
            # Processando o primeiro conjunto de dados (df1)
            df1 = pd.DataFrame(data1)
            for index1, row1 in df1.iterrows():
                str_topico1 = f"{row1['topicos']}, contexto: {str(row1['contexto'])}"
                conteudo_reuniao1 = docx_to_txt_diretorio(diretorio)
                if conteudo_reuniao1:
                    str_reuniao1 = json.dumps({'Reuniao': conteudo_reuniao1}, indent=4)  # Formatação com indentação
                    resultado_funcao1 = BPD_Transcricao(apiKey, nome_bpd, str_reuniao1, str_topico1, "")
    
                    # Armazenar o resultado no dicionário com o tópico como chave
                    resultados[row1['topicos']] = {
                        'contexto': row1['contexto'],
                        'resultado': resultado_funcao1
                    }
    
            # Processando o segundo conjunto de dados (df2)
            df2 = pd.DataFrame(data2)
    
            # Iterar conforme o número de processos em str_processo
            processos = json.loads(str_processo)
            for processo in processos:
                processo_nome = processo['processo']
    
                # Repetir duas vezes para cada processo: uma para Requerimento e outra para Detalhamento
                for index, row in df2.iterrows():
                    str_topico = f"{row['topicos']}, contexto: {str(row['contexto']).replace('[str_processo]', processo_nome)}"
                    conteudo_reuniao = docx_to_txt_diretorio(diretorio)
                    if conteudo_reuniao:
                        str_reuniao = json.dumps({'Reuniao': conteudo_reuniao}, indent=4)  # Formatação com indentação
                        resultado_funcao = BPD_Transcricao(apiKey, nome_bpd, str_reuniao, str_topico, "")
    
                        # Armazenar o resultado no dicionário com o tópico e processo como chave
                        topico_chave = f"{row['topicos']} do processo {processo_nome}"
                        resultados[topico_chave] = {
                            'contexto': row['contexto'].replace('[str_processo]', processo_nome),
                            'resultado': resultado_funcao
                        }
    
        except Exception as e:
            print(f"Erro ao processar a transição da reunião: {e}")
    
        return resultados  # Retornar o dicionário de resultados
    
    ################## str_processo
    # Filtrando as quatro colunas do DataFrame
    df_filtrado = df_bph[['processo']]
    
    # Removendo duplicatas
    df_filtrado_sem_duplicatas = df_filtrado.drop_duplicates()
    
    # Convertendo o DataFrame sem duplicatas para um objeto JSON
    str_processo = df_filtrado_sem_duplicatas.to_json(orient='records')
    
    # Convertendo o DataFrame sem duplicatas para uma lista de processos únicos
    processos_unicos = df_filtrado_sem_duplicatas['processo'].tolist()   
    
    # Chamada correta da função aninhada dentro de busca_variaveis
    ##### caminho_arquivo_reuniao = f'{caminho_base_bpd}\\{modulo}\\{nome_bpd}\\Transcricao\\' 
    caminho_arquivo_reuniao = f'{caminho_base_bpd}Transcricao//' 
    resultados_reuniao = transicao_reuniao(caminho_arquivo_reuniao, apiKey, nome_bpd, str_processo)
    
    # Convertendo o dicionário de resultados em uma string JSON formatada
    str_reuniao = json.dumps({'Reuniao': resultados_reuniao}, indent=4, ensure_ascii=False)  # ensure_ascii=False para manter caracteres especiais
    
    print(str_reuniao)  # Para visualização do resultado
        
    ###### str_doc   
    ##### diretorio_template_doc = f'{caminho_base_bpd}\\{modulo}\\{nome_bpd}\\Template\\'    
    
    
    diretorio_template_doc = f'{caminho_base_bpd}Template//'  
    str_doc = json.dumps({'Documento Modelo': docx_to_txt_diretorio(diretorio_template_doc)})
    
    ########################### str_nome_bpd_1
    # Filtrando as quatro colunas do DataFrame
    df_filtrado = df_bph[['nome_bpd']]
    
    # Removendo duplicatas
    df_filtrado_sem_duplicatas = df_filtrado.drop_duplicates()
    
    # Convertendo o DataFrame sem duplicatas para um objeto JSON
    str_nome_bpd_1 = df_filtrado_sem_duplicatas.to_json(orient='records')
     
    ################## str_modulo
    # Filtrando as quatro colunas do DataFrame
    df_filtrado = df_bph[['modulo']]
    
    # Removendo duplicatas
    df_filtrado_sem_duplicatas = df_filtrado.drop_duplicates()
    
    # Convertendo o DataFrame sem duplicatas para um objeto JSON
    str_modulo = df_filtrado_sem_duplicatas.to_json(orient='records')    
  
    ################## str_docs_scope_items
    # Filtrando as quatro colunas do DataFrame
    df_filtrado = df_bph[['scope_item','scope_items_docx']]
    
    # Removendo duplicatas
    df_filtrado_sem_duplicatas = df_filtrado.drop_duplicates()
    
    # Convertendo o DataFrame sem duplicatas para um objeto JSON
    str_docs_scope_items = df_filtrado_sem_duplicatas.to_json(orient='records')    

    ################## str_gap
    # Filtrando as quatro colunas do DataFrame
    df_filtrado = df_bph[['cod_requisito', 'requisito_descricao','cod_gap','gap_descricao','scope_item']]
    
    # Removendo duplicatas
    df_filtrado_sem_duplicatas = df_filtrado.drop_duplicates()
    
    # Convertendo o DataFrame sem duplicatas para um objeto JSON
    str_gaps = df_filtrado_sem_duplicatas.to_json(orient='records')       

    ################## str_nome_bpd
    # Filtrando as quatro colunas do DataFrame
    df_filtrado = df_bph[['nome_bpd']]
    
    # Removendo duplicatas
    df_filtrado_sem_duplicatas = df_filtrado.drop_duplicates()
    
    # Convertendo o DataFrame sem duplicatas para um objeto JSON
    str_nome_bpd = df_filtrado_sem_duplicatas.to_json(orient='records') 
    str_nome_bpd_html = gerar_h1(nome_bpd, 'Swis721 BT', '14', '#1F5CA9')  
    
    ################## str_idioma,str_empresa,str_segmento
    str_idioma = json.dumps({'Idioma': 'Portugues/br'})
    str_empresa = json.dumps({'Empresa': 'Potencial'})
    str_segmento = json.dumps({'Segmento': 'Oleo/gas'})
    
    ################## del
    del df_filtrado  
    del df_filtrado_sem_duplicatas
    
    return  str_idioma,str_empresa,str_segmento,str_processo,str_modulo,str_docs_scope_items,str_doc,str_reuniao,str_gaps,str_nome_bpd,str_nome_bpd_html

#%% DEF Prompts    
def BPD_01_01(apiKey, str_idioma, str_empresa, str_segmento, str_nome_bpd_1, str_processo, str_modulo, str_docs_scope_items, str_doc, str_reuniao):
    
    def def_BPD_01_01_result(apiKey, str_idioma, str_empresa, str_segmento, str_nome_bpd_1, str_processo, str_modulo, str_docs_scope_items, str_doc, str_reuniao):
        url = "https://gpt-templates.saiapplications.com"
        headers = {"X-Api-Key": apiKey}
        data = {
            "inputs": {
                "str_idioma": str_idioma,
                "str_empresa": str_empresa,
                "str_segmento": str_segmento,
                "str_nome_bpd": str_nome_bpd_1,
                "str_processo": str_processo,
                "str_modulo": str_modulo,
                "str_docs_scope_items": str_docs_scope_items,
                "str_doc": str_doc,
                "str_reuniao": str_reuniao,
            }
        }
        
        response = requests.post(f"{url}/api/templates/666213834c565b1488fb5424/execute", json=data, headers=headers)
        if response.status_code == 200:
            return response.text
        else:
            return f"Error: {response.status_code} - {response.text}"

    # Converte a string JSON em uma lista de dicionários
    dados = json.loads(str_processo)
    
    resultado = ""
    # Itera sobre cada dicionário na lista
    for item in dados:
        processo = item["processo"]
        resultado += def_BPD_01_01_result(apiKey, str_idioma, str_empresa, str_segmento, str_nome_bpd_1, processo, str_modulo, str_docs_scope_items, str_doc, str_reuniao)
        resultado += "\n\n"
       
    return resultado.rstrip("\n")
    

def BPD_01_1_01 (apiKey,str_idioma,str_nome_bpd_1,str_modulo,str_gaps):
    
    url = "https://gpt-templates.saiapplications.com"
    
    headers = {"X-Api-Key": apiKey}
    
    data = {    
    "inputs": {    
    "str_idioma": str_idioma,    
    "str_processo": str_nome_bpd_1,    
    "str_modulo": str_modulo,    
    "str_gaps": str_gaps,    
        }       
    }
    
    response = requests.post(f"{url}/api/templates/66312fbff577610990de6c2e/execute", json=data, headers=headers)   
    if response.status_code == 200:
        return(response.text)
  
def BPD_02_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
    "inputs": {
        "str_idioma ": str_idioma ,
        "str_contexto ": str_contexto ,
        "str_docs_scope_items": str_docs_scope_items,
        "str_doc": "",
        "str_gaps": "",
        "str_reuniao": str_reuniao,
    }
    }
    
    response = requests.post(f"{url}/api/templates/661949bddb0a64624fe3814f/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)
    
def BPD_03_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": "",
            "str_doc": str_doc,
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/66195858db0a64624fe381eb/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)


def BPD_04_01(apiKey, str_idioma, str_contexto, str_docs_scope_items, str_doc, str_gaps, str_reuniao, str_processo):
    
    def def_BPD_04_01_result(apiKey, str_idioma, str_contexto, str_docs_scope_items, str_doc, str_gaps, str_reuniao, processo):
        url = "https://gpt-templates.saiapplications.com"
        headers = {"X-Api-Key": apiKey}
        data = {
            "inputs": {
                "str_idioma": str_idioma,
                "str_contexto": str_contexto,
                "str_docs_scope_items": str_docs_scope_items,
                "str_doc": str_doc,
                "str_gaps": str_gaps,
                "str_reuniao": str_reuniao,
                "str_processo": processo,
            }
        }
        
        response = requests.post(f"{url}/api/templates/6663248e84660938876ccbd0/execute", json=data, headers=headers)
        if response.status_code == 200:
            return response.text
        else:
            return f"Erro ao executar template: {response.status_code} - {response.text}"

    # Converte a string JSON em uma lista de dicionários
    dados = json.loads(str_processo)
    
    resultado = ""
    # Itera sobre cada dicionário na lista
    for item in dados:
        processo = item["processo"]
        resultado += def_BPD_04_01_result(apiKey, str_idioma, str_contexto, str_docs_scope_items, str_doc, str_gaps, str_reuniao, processo)  
        resultado += "\n\n"
    
    return resultado.rstrip("\n")
       
    
def BPD_05_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": "",
            "str_doc": "",
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/661d533da3371a163fa1c9b7/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def BPD_06_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": "",
            "str_doc": str_doc,
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/661d5a05a3371a163fa1ca1b/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def BPD_07_01 (apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": "",
            "str_doc": str_doc,
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/661d5fe6db0a64624fe39ce4/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def BPD_08_01 (apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": "",
            "str_doc": str_doc,
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/661d644ea3371a163fa1ccc5/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)
    
def BPD_09_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": "",
            "str_doc": str_doc,
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/6626c48c6b4b6e690154be76/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)
 
def BPD_10_01 (apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": "",
            "str_doc": str_doc,
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/6626c5c96b4b6e690154be86/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def BPD_11_01 (apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": str_docs_scope_items,
            "str_doc": "",
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/6626cb7b6b4b6e690154bedb/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def BPD_12_01 (apiKey,str_idioma,str_nome_bpd_1,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_processo": str_nome_bpd_1,
            "str_gaps": str_gaps,
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/6626d4326b4b6e690154bf2a/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def BPD_13_01 (apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):

    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": str_docs_scope_items,
            "str_doc": "",
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/6626d44a6b4b6e690154bf2b/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

def BPD_14_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao):    
   
    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_contexto": str_contexto,
            "str_docs_scope_items": "",
            "str_doc": str_doc,
            "str_gaps": "",
            "str_reuniao": str_reuniao,
        }
    }
    
    response = requests.post(f"{url}/api/templates/6626d46f2b1023d8c1948e70/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)   

def BPD_Transcricao(apiKey,nome_bpd,str_reuniao,str_topico,str_processo):
    
    str_idioma = 'Portugues/BR'
    
    url = "https://gpt-templates.saiapplications.com"
    headers = {"X-Api-Key": apiKey}
    data = {
        "inputs": {
            "str_idioma": str_idioma,
            "str_reuniao": str_reuniao,
            "str_nome_bpd": nome_bpd,
            "str_processo": str_processo,
            "str_topico": str_topico,
        }
    }
    
    response = requests.post(f"{url}/api/templates/6661ee964c565b1488fb4f77/execute", json=data, headers=headers)
    if response.status_code == 200:
        return(response.text)

#%% format_html_in_docxfrom bs4 import BeautifulSoup
def clear_paragraph(paragraph):
    """Limpa o texto de um parágrafo no documento."""
    p_element = paragraph._element
    p_child_elements = [elm for elm in p_element.iterchildren()]
    for child_element in p_child_elements:
        p_element.remove(child_element)

def set_cell_borders(cell, border_color='000000', border_sz=4):
    """Função para configurar as bordas de uma célula."""
    directions = ('top', 'left', 'bottom', 'right')
    for direction in directions:
        border_elm = OxmlElement('w:tcBorders')
        border_sub_elm = OxmlElement(f'w:{direction}')
        border_sub_elm.set(qn('w:val'), 'single')
        border_sub_elm.set(qn('w:sz'), str(border_sz))  # Espessura da borda
        border_sub_elm.set(qn('w:space'), '0')
        border_sub_elm.set(qn('w:color'), border_color)
        border_elm.append(border_sub_elm)
        cell._element.get_or_add_tcPr().append(border_elm)

def apply_styles(run, style_str):
    """Aplica estilos CSS a um run do Word."""
    styles = style_str.split(';')
    for style in styles:
        if style.strip():
            prop, value = style.split(':')
            prop = prop.strip()
            value = value.strip()
            if prop == 'font-family':
                run.font.name = value
            elif prop == 'font-size':
                run.font.size = Pt(int(value.replace('px', '')))
            elif prop == 'color':
                color_value = value.lstrip('#')
                run.font.color.rgb = RGBColor(int(color_value[0:2], 16), int(color_value[2:4], 16), int(color_value[4:6], 16))

def insert_formatted_content_from_html(paragraph, html_content):
    """Insere conteúdo formatado de HTML em um parágrafo do documento."""
    soup = BeautifulSoup(html_content, 'html.parser')

    def add_text(paragraph, node):
        if node.name in ['p', 'h1', 'li']:
            process_special_tags(paragraph, node)
        elif node.name == 'table':
            process_table(paragraph, node)
        elif node.name == 'strong':
            run = paragraph.add_run(node.get_text())
            run.bold = True
        elif node.name == 'em':
            run = paragraph.add_run(node.get_text())
            run.italic = True
        elif node.name == 'u':
            run = paragraph.add_run(node.get_text())
            run.underline = True
        elif node.name == 'code':
            code_text = node.get_text(separator="\n")
            code_lines = code_text.split('\n')
            for line in code_lines:
                run = paragraph.add_run(line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 69, 0)
                paragraph.add_run('\n')
        else:
            if node.string:
                run = paragraph.add_run(node.string)
                if node.parent.name == 'span' and 'style' in node.parent.attrs:
                    apply_styles(run, node.parent['style'])
            else:
                for child in node.children:
                    add_text(paragraph, child)

    def process_special_tags(paragraph, element):
        if element.name == 'p':
            for node in element:
                add_text(paragraph, node)
            paragraph.add_run('\n')
        elif element.name == 'h1':
            for child in element.children:
                if child.name == 'span' and 'style' in child.attrs:
                    run = paragraph.add_run(child.get_text())
                    apply_styles(run, child['style'])
                else:
                    add_text(paragraph, child)
            paragraph.add_run('\n')
        elif element.name == 'li':
            run = paragraph.add_run('• ')
            if 'style' in element.attrs:
                apply_styles(run, element['style'])
    
            for child in element.children:
                add_text(paragraph, child)
    
            paragraph.add_run('\n')

    def process_table(paragraph, table_element):
        doc = paragraph.part.document
        num_cols = max(len(row.find_all(['td', 'th'])) for row in table_element.find_all('tr'))
        table = doc.add_table(rows=0, cols=num_cols)
        
        col_width = Inches(1.0 / num_cols)

        for row_element in table_element.find_all('tr'):
            row = table.add_row()
            row_cells = row_element.find_all(['td', 'th'])
            for i, cell_element in enumerate(row_cells):
                cell = row.cells[i]
                cell.width = col_width
                clear_paragraph(cell.paragraphs[0])
                add_text(cell.paragraphs[0], cell_element)
                set_cell_borders(cell)

                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                cell._element.get_or_add_tcPr().append(shading_elm)
                
                if cell_element.name == 'th':
                    run.font.bold = True
                    run.font.size = Pt(12)

        return table

    body = soup.find('body')
    if body:
        for node in body.children:
            if node.name == 'table':
                new_paragraph = paragraph.insert_paragraph_before()
                table = process_table(new_paragraph, node)
                paragraph._element.addnext(table._element)
            else:
                add_text(paragraph, node)

def process_docx_with_html(file_path):
    """Abre um documento Word, formata o HTML contido e salva o resultado."""
    
    #file_path = r'C:\temp\Vendas 5.docx'
    doc = Document(file_path)
    html_pattern = re.compile(r'<!DOCTYPE html>.*?</html>', re.DOTALL)

    for paragraph in doc.paragraphs:
        matches = html_pattern.findall(paragraph.text)
        if matches:
            for i, html_content in enumerate(matches):
                clear_paragraph(paragraph)
                insert_formatted_content_from_html(paragraph, html_content)
                if i < len(matches) - 1:
                    paragraph = paragraph.insert_paragraph_before()
    
    #formatted_path = file_path.replace('.docx', '_formatted.docx')
    formatted_path = file_path.replace('.docx', '.docx')
    
    doc.save(formatted_path)
    print(f'Arquivo salvo com sucesso em: {formatted_path}')
    
#%%  replace_word
def substituir_no_word(modelo_docx, df, destino_docx):
    doc = Document(modelo_docx)
    for paragraph in doc.paragraphs:
        for index, row in df.iterrows():
            if row['replace_word'] in paragraph.text:
                str_resultado = str(row['html_content']) if pd.notna(row['html_content']) else ''
                paragraph.text = paragraph.text.replace(row['replace_word'], str_resultado)
    doc.save(destino_docx)
    # Chama a função para formatar o HTML após a substituição
    process_docx_with_html(destino_docx)

#%% Def Main()
def main():   
   
    ############################################## Caminho do bpd
    caminho_base_bpd = '/caminho/ficticio/para/portfolio/'    
        
    ############################################## Caminho do arquivo Excel
    nome_arquivo_xlsx = 'planilha_ficticia.xlsx'
    
    ############################################## Caminho para o diretório dos Scope Itens
    caminho_scope_itens = '/caminho/ficticio/para/portfolio/ScopeItems/'

    ############################################## Caminho para o diretório dos Scope Itens
    caminho_modelo_docx_bpd = '/caminho/ficticio/para/portfolio/'
    
    
    
    #df_resultado = cria_df_BPH(caminho_base_bpd, nome_arquivo_xlsx, caminho_scope_itens)
    #print(df_resultado)
    
    # Cria df_bph 
    df_bph = cria_df_BPH(caminho_base_bpd,nome_arquivo_xlsx,caminho_scope_itens)
    
    #df_bph = df_bph[df_bph['nome_bpd'] == 'Contas a Pagar']

    dados_unicos = df_bph[['nome_bpd', 'modulo']].drop_duplicates()
    
    # Converter para uma lista de tuplas
    dados_unicos = [tuple(x) for x in dados_unicos.to_numpy()]
    
    

    # Loop sobre os valores únicos e filtrando a DataFrame para cada valor único
    for nome_bpd, modulo in dados_unicos:
        print(nome_bpd)

        # str_caminho = f'{caminho_base_bpd}{modulo}\\{nome_bpd}\\'     
        # destino_docx = f'{caminho_base_bpd}{modulo}\\\\{nome_bpd}\\{nome_bpd}.docx'        
        # modelo_docx = f'{caminho_base_bpd}modelo_potencial.docx'
         
        nome_bpd_arquivo = replace_especial(nome_bpd)
        
        destino_docx = f'{caminho_base_bpd}\\{nome_bpd_arquivo}.docx'        
        modelo_docx = f'{caminho_modelo_docx_bpd}modelo_potencial.docx'              
        
        # Filtrando a DataFrame base
        df_filtered = df_bph[df_bph['nome_bpd'] == nome_bpd]
        str_idioma, str_empresa, str_segmento, str_processo, str_modulo, str_docs_scope_items, str_doc, str_reuniao, str_gaps, str_nome_bpd,return_nome_bpd = busca_variaveis(df_filtered, nome_bpd, apiKey, caminho_base_bpd, modulo)

       
        ################## PROMPT's         
                 
        dt_resultPrompts = pd.DataFrame(columns=['replace_word', 'html_content'])
        i = 0
                  
        print('return_BPD_01_01')
        
        return_BPD_01_01 = BPD_01_01(apiKey, str_idioma, str_empresa, str_segmento, str_nome_bpd, str_processo, str_modulo, str_docs_scope_items, str_doc, str_reuniao)
        
        dt_resultPrompts.loc[i] = ['@BPD_01_01',return_BPD_01_01]
        
        print ('return_BPD_01_1_01')
        return_BPD_01_1_01 = BPD_01_1_01(apiKey,str_idioma,str_nome_bpd,str_modulo,str_gaps)
        return_BPD_01_1_01 = return_BPD_01_1_01.replace('null', '')  
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_01_1_01', return_BPD_01_1_01]
    
        print ('return_BPD_02_01')    
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}'    
        return_BPD_02_01 = BPD_02_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)
        i=i+1
        dt_resultPrompts.loc[i]= ['@BPD_02_01', return_BPD_02_01]
   
        print ('return_BPD_03_01')    
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}'
        return_BPD_03_01 = BPD_03_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)   
        return_BPD_03_01 = return_BPD_03_01.replace('null', '')
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_03_01', return_BPD_03_01]
        
        #  dt_resultPrompts=dt_resultPrompts.drop(4)
    
        print('return_BPD_04_01')
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}'
        return_BPD_04_01 = BPD_04_01(apiKey, str_idioma, str_contexto, str_docs_scope_items, str_doc, str_gaps, str_reuniao, str_processo)
        return_BPD_04_01 = return_BPD_04_01.replace('null', '')
        i = i + 1
        dt_resultPrompts.loc[i] = ['@BPD_04_01', return_BPD_04_01]               
        
    
        print ('return_BPD_05_01')      
        return_BPD_05_01 = BPD_05_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)    
        return_BPD_05_01 = return_BPD_05_01.replace('null', '')
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_05_01', return_BPD_05_01]
    
        print ('return_BPD_06_01')       
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}{return_BPD_04_01}{return_BPD_05_01}'
        return_BPD_06_01 = BPD_06_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)    
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_06_01', return_BPD_06_01]
    
        print ('return_BPD_07_01')       
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}{return_BPD_04_01}{return_BPD_05_01}{return_BPD_06_01}'
        return_BPD_07_01 = BPD_07_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)        
        return_BPD_07_01 = return_BPD_07_01.replace('null', '')
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_07_01', return_BPD_07_01]
    
        print ('return_BPD_08_01')      
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}{return_BPD_04_01}{return_BPD_05_01}{return_BPD_06_01}{return_BPD_07_01}'
        return_BPD_08_01 = BPD_08_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)      
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_08_01', return_BPD_08_01]
    
        print ('return_BPD_09_01')           
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}{return_BPD_04_01}{return_BPD_05_01}{return_BPD_06_01}{return_BPD_07_01}{return_BPD_08_01}'
        return_BPD_09_01 = BPD_09_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)       
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_09_01', return_BPD_09_01]
    
        print ('return_BPD_10_01')     
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}{return_BPD_04_01}{return_BPD_05_01}{return_BPD_06_01}{return_BPD_07_01}{return_BPD_08_01},{return_BPD_09_01}'
        return_BPD_10_01 = BPD_10_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)      
        return_BPD_10_01 = return_BPD_10_01.replace('null', '')
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_10_01', return_BPD_10_01]
    
        print ('return_BPD_11_01')      
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}{return_BPD_04_01}{return_BPD_05_01}{return_BPD_06_01}{return_BPD_07_01}{return_BPD_08_01},{return_BPD_09_01}{return_BPD_10_01}'
        return_BPD_11_01 = BPD_11_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)         
        return_BPD_11_01 = return_BPD_11_01.replace('null', '')
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_11_01', return_BPD_11_01]
    
        print ('return_BPD_12_01')      
        return_BPD_12_01 = BPD_12_01(apiKey,str_idioma,str_processo,str_gaps,str_reuniao)  
        return_BPD_12_01 = return_BPD_12_01.replace('null', '')
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_12_01', return_BPD_12_01]
    
        print ('return_BPD_13_01')      
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}{return_BPD_04_01}{return_BPD_05_01}{return_BPD_06_01}{return_BPD_07_01}{return_BPD_08_01},{return_BPD_09_01}{return_BPD_10_01}{return_BPD_11_01},{return_BPD_12_01}'
        return_BPD_13_01 = BPD_13_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)  
        return_BPD_13_01 = return_BPD_13_01.replace('null', '')
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_13_01', return_BPD_13_01]
        
        print ('return_BPD_14_01')      
        str_contexto = f'{return_BPD_01_01}{return_BPD_01_1_01}{return_BPD_02_01}{return_BPD_03_01}{return_BPD_04_01}{return_BPD_05_01}{return_BPD_06_01}{return_BPD_07_01}{return_BPD_08_01},{return_BPD_09_01}{return_BPD_10_01}{return_BPD_11_01},{return_BPD_12_01},{return_BPD_13_01}'
        return_BPD_14_01 = BPD_14_01(apiKey,str_idioma,str_contexto,str_docs_scope_items,str_doc,str_gaps,str_reuniao)
        return_BPD_14_01 = return_BPD_14_01.replace('null', '')    
        i=i+1
        dt_resultPrompts.loc[i] = ['@BPD_14_01', return_BPD_14_01] 
        
        i=i+1
        dt_resultPrompts.loc[i] = ['@NOME_BPD', return_nome_bpd]            
        
    ###################### replace_word
        
        substituir_no_word(modelo_docx, dt_resultPrompts, destino_docx)
        
        
if __name__ == "__main__":
    ## main()